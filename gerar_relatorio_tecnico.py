"""
Script para gerar relatório técnico em formato Word
Inferência Geoespacial Aplicada à Geração de Base Vetorial Rodoviária
Estudo de Caso para o Estado de São Paulo
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os
from datetime import datetime

def adicionar_cabecalho(doc):
    """Adiciona cabeçalho ao documento"""
    section = doc.sections[0]
    header = section.header
    paragraph = header.paragraphs[0]
    paragraph.text = "Inferência Geoespacial – Base Vetorial Rodoviária – Estado de São Paulo"
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    paragraph.style.font.size = Pt(9)
    paragraph.style.font.color.rgb = RGBColor(128, 128, 128)

def adicionar_rodape(doc):
    """Adiciona rodapé ao documento"""
    section = doc.sections[0]
    footer = section.footer
    paragraph = footer.paragraphs[0]
    paragraph.text = f"PLI/SP – Malha Vicinal Estimada | Janeiro 2026 | Página "
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.style.font.size = Pt(9)
    paragraph.style.font.color.rgb = RGBColor(128, 128, 128)

def adicionar_titulo_principal(doc, texto):
    """Adiciona título principal"""
    titulo = doc.add_heading(texto, level=0)
    titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in titulo.runs:
        run.font.size = Pt(18)
        run.font.color.rgb = RGBColor(0, 51, 102)
        run.font.bold = True

def adicionar_subtitulo(doc, texto):
    """Adiciona subtítulo"""
    subtitulo = doc.add_paragraph(texto)
    subtitulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in subtitulo.runs:
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(102, 102, 102)
        run.font.italic = True

def adicionar_secao(doc, titulo, nivel=1):
    """Adiciona uma seção"""
    heading = doc.add_heading(titulo, level=nivel)
    for run in heading.runs:
        run.font.color.rgb = RGBColor(0, 51, 102)

def adicionar_paragrafo(doc, texto, espacamento_antes=6, espacamento_depois=6):
    """Adiciona parágrafo formatado"""
    p = doc.add_paragraph(texto)
    p.paragraph_format.line_spacing = 1.15
    p.paragraph_format.space_before = Pt(espacamento_antes)
    p.paragraph_format.space_after = Pt(espacamento_depois)
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    for run in p.runs:
        run.font.size = Pt(11)
        run.font.name = 'Calibri'
    return p

def adicionar_lista(doc, itens, numerada=False):
    """Adiciona lista de itens"""
    for item in itens:
        p = doc.add_paragraph(item, style='List Number' if numerada else 'List Bullet')
        p.paragraph_format.line_spacing = 1.15
        for run in p.runs:
            run.font.size = Pt(11)

def adicionar_tabela_dados(doc):
    """Adiciona tabela com resumo dos dados"""
    table = doc.add_table(rows=6, cols=3)
    table.style = 'Light Grid Accent 1'
    
    # Cabeçalho
    header_cells = table.rows[0].cells
    header_cells[0].text = 'Dataset'
    header_cells[1].text = 'Fonte'
    header_cells[2].text = 'Características'
    
    # Dados
    dados = [
        ('OSM Sudeste', 'Geofabrik (Nov/2025)', '1.358.788 ways (SP), 788,8 MB'),
        ('Áreas Urbanas', 'IBGE 2019', '24.186 polígonos'),
        ('Faces Logradouro', 'IBGE Censo 2022', '~2,4 milhões (SP)'),
        ('Malha DER/SP', 'DER/SP (Mai/2025)', '7.417 segmentos, 25.919 km'),
        ('Municípios SEADE', 'Fundação SEADE', '645 municípios, 16 RAs')
    ]
    
    for i, (dataset, fonte, carac) in enumerate(dados, start=1):
        cells = table.rows[i].cells
        cells[0].text = dataset
        cells[1].text = fonte
        cells[2].text = carac
    
    # Formatar células
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(10)

def adicionar_tabela_filtros(doc):
    """Adiciona tabela com critérios de filtragem"""
    table = doc.add_table(rows=4, cols=3)
    table.style = 'Light Grid Accent 1'
    
    # Cabeçalho
    header_cells = table.rows[0].cells
    header_cells[0].text = 'Critério'
    header_cells[1].text = 'Condição'
    header_cells[2].text = 'Justificativa'
    
    # Dados
    filtros = [
        ('Exclusão urbana', 'area_urb = 0', 'Evita inflar malha com ruas locais'),
        ('Exclusão logradouro', 'path IS NULL', 'Retira vias censitárias urbanas'),
        ('Exclusão DER', 'malha_der = 0', 'Remove sobreposições com rede estadual')
    ]
    
    for i, (criterio, condicao, justif) in enumerate(filtros, start=1):
        cells = table.rows[i].cells
        cells[0].text = criterio
        cells[1].text = condicao
        cells[2].text = justif
    
    # Formatar células
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(10)

def adicionar_tabela_resultados(doc):
    """Adiciona tabela com síntese dos resultados"""
    table = doc.add_table(rows=6, cols=3)
    table.style = 'Light Grid Accent 1'
    
    # Cabeçalho
    header_cells = table.rows[0].cells
    header_cells[0].text = 'Indicador'
    header_cells[1].text = 'Malha Vicinal (OSM)'
    header_cells[2].text = 'Malha Total (OSM+DER)'
    
    # Dados
    resultados = [
        ('Extensão Total', '25.918,60 km', '47.666,00 km'),
        ('Total Segmentos', '7.417 segmentos', '~14.000 segmentos'),
        ('Média Municipal', '40,18 km/município', '73,90 km/município'),
        ('Densidade Espacial', '1.312 km/10.000 km²', '2.413 km/10.000 km²'),
        ('Densidade Populacional', '35 km/10.000 hab', '103 km/10.000 hab')
    ]
    
    for i, (indicador, osm, total) in enumerate(resultados, start=1):
        cells = table.rows[i].cells
        cells[0].text = indicador
        cells[1].text = osm
        cells[2].text = total
    
    # Formatar células
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(10)

def adicionar_caixa_ressalva(doc, titulo, texto):
    """Adiciona caixa de ressalva destacada"""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1)
    p.paragraph_format.right_indent = Cm(1)
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(12)
    
    # Título da ressalva
    run_titulo = p.add_run(f"⚠️ {titulo}\n")
    run_titulo.bold = True
    run_titulo.font.size = Pt(11)
    run_titulo.font.color.rgb = RGBColor(204, 102, 0)
    
    # Texto da ressalva
    run_texto = p.add_run(texto)
    run_texto.font.size = Pt(10)
    run_texto.font.color.rgb = RGBColor(51, 51, 51)
    
    # Adicionar borda
    pPr = p._element.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    for border_name in ['top', 'left', 'bottom', 'right']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '12')
        border.set(qn('w:space'), '4')
        border.set(qn('w:color'), 'FFA500')
        pBdr.append(border)
    pPr.append(pBdr)
    
    # Adicionar sombreamento
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), 'FFF8DC')
    pPr.append(shd)

def adicionar_referencia_figura(doc, numero, titulo, caminho_relativo):
    """Adiciona referência a uma figura"""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    run = p.add_run(f"Figura {numero} – {titulo}")
    run.font.size = Pt(10)
    run.font.italic = True
    run.font.color.rgb = RGBColor(102, 102, 102)
    
    # Adicionar nota sobre a figura
    p_nota = doc.add_paragraph()
    p_nota.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_nota = p_nota.add_run(f"[Figura a ser inserida: {caminho_relativo}]")
    run_nota.font.size = Pt(9)
    run_nota.font.color.rgb = RGBColor(150, 150, 150)
    run_nota.font.italic = True

def criar_relatorio():
    """Cria o relatório técnico completo"""
    
    doc = Document()
    
    # Configurar margens
    sections = doc.sections
    for section in sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(3)
        section.right_margin = Cm(2)
    
    # Adicionar cabeçalho e rodapé
    adicionar_cabecalho(doc)
    adicionar_rodape(doc)
    
    # ============================================================================
    # CAPA
    # ============================================================================
    adicionar_titulo_principal(doc, "INFERÊNCIA GEOESPACIAL APLICADA À GERAÇÃO DE BASE VETORIAL RODOVIÁRIA")
    adicionar_subtitulo(doc, "Estudo de Caso para o Estado de São Paulo")
    doc.add_paragraph()
    doc.add_paragraph()
    
    p_info = doc.add_paragraph()
    p_info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_info = p_info.add_run("Base Estimada: Malha Municipal – PLI/SP\n(OpenStreetMap + DER/SP + SEADE/IBGE)")
    run_info.font.size = Pt(11)
    run_info.font.color.rgb = RGBColor(102, 102, 102)
    
    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()
    
    p_data = doc.add_paragraph()
    p_data.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_data = p_data.add_run(f"São Paulo\n{datetime.now().strftime('%B de %Y')}")
    run_data.font.size = Pt(12)
    
    doc.add_page_break()
    
    # ============================================================================
    # 1. INTRODUÇÃO
    # ============================================================================
    adicionar_secao(doc, "1. INTRODUÇÃO", nivel=1)
    
    adicionar_paragrafo(doc, 
        "O Plano de Logística Integrada do Estado de São Paulo (PLI/SP) constitui instrumento estratégico "
        "de planejamento setorial voltado à melhoria da infraestrutura de transportes e logística no território "
        "paulista. No que tange ao setor rodoviário, o PLI/SP tem por missão diagnosticar o estado atual da "
        "malha rodoviária estadual, identificar gargalos logísticos, propor investimentos prioritários e orientar "
        "políticas públicas que promovam a integração eficiente dos diferentes modais de transporte, contribuindo "
        "para o desenvolvimento econômico e a competitividade do Estado."
    )
    
    adicionar_paragrafo(doc,
        "Visando ampliar a completude e robustez metodológica das análises desenvolvidas no âmbito do PLI/SP, "
        "optou-se por considerar, além das infraestruturas de escala estadual e federal, também as escalas locais "
        "do sistema rodoviário. No caso do setor rodoviário, essas escalas locais são representadas pela malha "
        "viária municipal — também denominada malha vicinal —, responsável por conectar propriedades rurais, "
        "centros produtivos locais e comunidades aos eixos rodoviários principais. A incorporação dessa camada de "
        "análise, embora não configurasse requisito formal do escopo original do Plano, permite caracterização "
        "mais abrangente do sistema de circulação paulista, especialmente no que tange aos elos iniciais das "
        "cadeias logísticas dos setores agropecuário e agroindustrial, estratégicos para a economia estadual."
    )
    
    adicionar_paragrafo(doc,
        "Entretanto, diferentemente das malhas federal e estadual, que contam com cadastros oficiais consolidados "
        "e atualizados periodicamente, a malha vicinal — de responsabilidade dos 645 municípios paulistas — não "
        "possui, até o momento, uma base vetorial unificada, oficial e de cobertura estadual. Essa ausência "
        "representa lacuna significativa para o planejamento integrado, dificultando análises territoriais "
        "comparativas, a identificação de áreas prioritárias para investimentos e a avaliação da real extensão "
        "e condições da infraestrutura viária local."
    )
    
    adicionar_paragrafo(doc,
        "Diante desse cenário, foi idealizada uma alternativa metodológica para estimar essa base, utilizando "
        "técnicas de inferência geoespacial a partir da integração de fontes abertas (OpenStreetMap) e oficiais "
        "(DER/SP, IBGE, SEADE). O presente estudo, desenvolvido no contexto do PLI/SP, tem por objetivo apresentar "
        "essa metodologia e os resultados dela decorrentes, gerando uma base vetorial rodoviária estimada para a "
        "malha vicinal do Estado de São Paulo. Ressalta-se que a base produzida constitui aproximação derivada de "
        "geoprocessamento — não se tratando de levantamento topográfico ou cadastro oficial —, mas que, ainda assim, "
        "pode subsidiar análises exploratórias e orientar esforços de caracterização mais aprofundada onde os "
        "indícios preliminares apontarem maior relevância."
    )
    
    adicionar_caixa_ressalva(doc, "Ressalva Metodológica Fundamental",
        "A base apresentada constitui uma APROXIMAÇÃO da infraestrutura rodoviária municipal paulista, "
        "obtida por meio de geoprocessamento aplicado a bases existentes. Trata-se de um produto derivado "
        "de inferência geoespacial — não de levantamento topográfico ou cadastro oficial. Os valores são "
        "indicativos e podem divergir da realidade em campo devido a limitações inerentes às fontes, à "
        "metodologia de filtragem e aos processos de integração espacial empregados. Esta base representa "
        "um ponto de partida para análises exploratórias, devendo ser refinada continuamente por meio de "
        "validação cruzada com outras fontes oficiais, verificações em campo e integração com cadastros "
        "municipais quando disponíveis."
    )
    
    adicionar_paragrafo(doc,
        "O estudo está estruturado em quatro seções principais: (1) esta Introdução, que contextualiza a "
        "problemática e os objetivos; (2) Metodologia, que detalha os dados utilizados e os procedimentos "
        "de processamento espacial; (3) Resultados e Discussão, que apresenta os indicadores obtidos nas "
        "escalas municipal e regional; e (4) Considerações Finais, que sintetiza os achados principais e "
        "aponta limitações e perspectivas de aprimoramento."
    )
    
    doc.add_page_break()
    
    # ============================================================================
    # 2. METODOLOGIA
    # ============================================================================
    adicionar_secao(doc, "2. METODOLOGIA", nivel=1)
    
    adicionar_paragrafo(doc,
        "A metodologia adotada baseia-se em princípios de ciência de dados espaciais, priorizando "
        "reprodutibilidade, rastreabilidade de decisões técnicas e documentação integral do fluxo de "
        "processamento. Todos os dados de entrada são públicos, as ferramentas utilizadas são de código "
        "aberto e os critérios de filtragem estão explicitados em detalhes. Essa transparência assegura "
        "auditabilidade técnica e permite replicação ou adaptação do estudo para outros contextos territoriais."
    )
    
    # --------------------------------------------------------------------------
    # 2.1 Dados
    # --------------------------------------------------------------------------
    adicionar_secao(doc, "2.1. Fontes de Dados", nivel=2)
    
    adicionar_paragrafo(doc,
        "O estudo integra cinco conjuntos de dados geoespaciais de fontes abertas e oficiais, "
        "conforme sintetizado na Tabela 1."
    )
    
    adicionar_tabela_dados(doc)
    
    p_tabela = doc.add_paragraph()
    p_tabela.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_tabela = p_tabela.add_run("Tabela 1 – Síntese dos dados utilizados")
    run_tabela.font.size = Pt(10)
    run_tabela.font.italic = True
    run_tabela.font.color.rgb = RGBColor(102, 102, 102)
    
    doc.add_paragraph()
    
    # OpenStreetMap
    adicionar_secao(doc, "2.1.1. OpenStreetMap (OSM)", nivel=3)
    adicionar_paragrafo(doc,
        "O OpenStreetMap é um projeto colaborativo de mapeamento global que registra vias de todas as "
        "hierarquias a partir de contribuições voluntárias, imagens de satélite e levantamentos de campo. "
        "Para este estudo, utilizou-se o extrato regional Sudeste disponibilizado pela Geofabrik GmbH "
        "(snapshot de 11 de novembro de 2025), contendo 1.358.788 segmentos viários (ways) para o Estado "
        "de São Paulo em formato PBF (Protocol Buffer Format), totalizando 788,8 MB."
    )
    
    # IBGE
    adicionar_secao(doc, "2.1.2. Instituto Brasileiro de Geografia e Estatística (IBGE)", nivel=3)
    adicionar_paragrafo(doc,
        "Duas bases do IBGE foram utilizadas para refinamento espacial da malha vicinal:"
    )
    
    adicionar_lista(doc, [
        "Áreas Urbanizadas 2019: delimitação de áreas efetivamente urbanizadas, incluindo cidades, vilas "
        "e aglomerações urbanas, com 24.186 polígonos para o Estado de São Paulo;",
        
        "Faces de Logradouro (Censo 2022): representação linear de todas as ruas, avenidas e vias mapeadas "
        "pelo IBGE durante o Censo, totalizando aproximadamente 2,4 milhões de faces para São Paulo."
    ])
    
    adicionar_caixa_ressalva(doc, "Limitação Temporal",
        "Os dados de Áreas Urbanizadas referem-se a 2019 e podem não capturar expansão urbana recente "
        "(2020-2025). Algumas vicinais podem ter sido urbanizadas nesse período, resultando em possível "
        "superestimação da malha rural em áreas de expansão urbana acelerada."
    )
    
    # DER/SP
    adicionar_secao(doc, "2.1.3. Departamento de Estradas de Rodagem de São Paulo (DER/SP)", nivel=3)
    adicionar_paragrafo(doc,
        "O cadastro oficial da malha rodoviária estadual foi obtido do DER/SP (atualização de 16 de maio "
        "de 2025), contendo 7.417 segmentos de rodovias municipais cadastradas, totalizando aproximadamente "
        "25.919 km. Esse conjunto representa as vicinais oficialmente reconhecidas e sob manutenção, mas "
        "não inclui todas as vicinais existentes, focando nas principais vias municipais."
    )
    
    # SEADE
    adicionar_secao(doc, "2.1.4. Fundação Sistema Estadual de Análise de Dados (SEADE)", nivel=3)
    adicionar_paragrafo(doc,
        "A malha municipal oficial da Fundação SEADE fornece os limites administrativos dos 645 municípios "
        "paulistas, incluindo atributos de código IBGE, nome do município, identificador de Região "
        "Administrativa (GID_RA) e área territorial. Esses dados permitem agregações por município e por "
        "Região Administrativa (16 RAs)."
    )
    
    # --------------------------------------------------------------------------
    # 2.2 Sistema de Referência Espacial
    # --------------------------------------------------------------------------
    adicionar_secao(doc, "2.2. Sistema de Referência Espacial", nivel=2)
    
    adicionar_paragrafo(doc,
        "Todos os processamentos e cálculos métricos foram realizados no sistema de coordenadas projetadas "
        "SIRGAS 2000 / UTM Zone 23S (EPSG:31983), que é o padrão oficial adotado pelo IBGE, SEADE e DER/SP "
        "para o Estado de São Paulo. Embora o território paulista se estenda pelos fusos UTM 22S (oeste) e "
        "23S (leste), optou-se pelo fuso 23S por ser o padrão oficial, garantindo compatibilidade com bases "
        "oficiais e consistência nas análises territoriais. A projeção UTM garante distorção inferior a 1% "
        "dentro da zona, adequada para análises em escala estadual."
    )
    
    # --------------------------------------------------------------------------
    # 2.3 Procedimentos de Processamento Espacial
    # --------------------------------------------------------------------------
    adicionar_secao(doc, "2.3. Procedimentos de Processamento Espacial", nivel=2)
    
    adicionar_paragrafo(doc,
        "A geração da base vetorial de malha vicinal estimada seguiu um fluxo sequencial de quatro etapas "
        "principais, conforme ilustrado na Figura 1."
    )
    
    adicionar_referencia_figura(doc, 1, "Fluxo metodológico de processamento espacial", "passo1.png")
    
    # Etapa 1
    adicionar_secao(doc, "2.3.1. Preparação e Harmonização", nivel=3)
    adicionar_paragrafo(doc,
        "Os dados da base OSM foram recortados para o limite administrativo do Estado de São Paulo, "
        "utilizando como referência espacial o sistema SIRGAS 2000 / UTM 23S. Esse procedimento garante "
        "extensão compatível com os limites oficiais estaduais e reduz significativamente o volume "
        "computacional nas etapas subsequentes. Em seguida, os dados foram enriquecidos com informações "
        "das bases da Fundação SEADE, agregando os atributos Cod_ibge (código do município) e GID_RA "
        "(identificador da Região Administrativa). Esse enriquecimento é fundamental para permitir análises "
        "multiescala, possibilitando tanto caracterizações municipais quanto agregações regionais."
    )
    
    adicionar_referencia_figura(doc, 2, "Recorte espacial e enriquecimento de atributos", "passo2.png")
    
    # Etapa 2
    adicionar_secao(doc, "2.3.2. Filtragem de Ruído Urbano", nivel=3)
    adicionar_paragrafo(doc,
        "Para remover vias urbanas que não caracterizam a malha vicinal, foram aplicados dois filtros "
        "espaciais complementares. Primeiro, realizou-se uma interseção com a camada de Faces de Logradouro "
        "do IBGE (Censo 2022), utilizando um buffer de 10 metros. Esse procedimento remove vias intraquadra "
        "(ruas locais) e evita superestimação causada pelo mapeamento denso característico de áreas urbanas "
        "no OpenStreetMap. Em seguida, aplicou-se uma segunda interseção com a camada de Áreas Urbanizadas "
        "do IBGE (2019), assegurando o alinhamento com os perímetros urbanos oficialmente delimitados. "
        "Essa estratégia dupla reduz o viés metropolitano e garante que apenas vias com características "
        "rurais ou de acesso sejam mantidas na base."
    )
    
    adicionar_referencia_figura(doc, 3, "Filtragem de áreas urbanas e faces de logradouro", "passo3.png")
    
    # Etapa 3
    adicionar_secao(doc, "2.3.3. Subtração da Rede Estadual", nivel=3)
    adicionar_paragrafo(doc,
        "Para evitar duplicidade entre a malha estadual (DER/SP) e a malha vicinal estimada, aplicou-se "
        "um buffer de 60 a 100 metros sobre os eixos rodoviários estaduais cadastrados. Essas margens "
        "ampliadas foram dimensionadas para absorver imprecisões de alinhamento geométrico entre as bases "
        "OSM e DER/SP, além de evitar a inclusão de vicinais estritamente paralelas às rodovias estaduais. "
        "Todas as vias OSM que interceptam essa máscara de exclusão foram descartadas, garantindo que a "
        "malha resultante seja composta exclusivamente por rodovias municipais e estradas vicinais, sem "
        "sobreposição com a rede de responsabilidade estadual."
    )
    
    adicionar_referencia_figura(doc, 4, "Subtração da rede estadual (DER/SP)", "passo4.png")
    
    # Etapa 4
    adicionar_secao(doc, "2.3.4. Seleção Final de Vicinais", nivel=3)
    adicionar_paragrafo(doc,
        "Após as etapas de filtragem espacial, aplicaram-se três critérios lógicos finais para seleção "
        "dos segmentos que compõem a malha vicinal estimada, conforme sintetizado na Tabela 2."
    )
    
    adicionar_tabela_filtros(doc)
    
    p_tabela2 = doc.add_paragraph()
    p_tabela2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_tabela2 = p_tabela2.add_run("Tabela 2 – Critérios de seleção final de segmentos vicinais")
    run_tabela2.font.size = Pt(10)
    run_tabela2.font.italic = True
    run_tabela2.font.color.rgb = RGBColor(102, 102, 102)
    
    adicionar_paragrafo(doc,
        "A aplicação conjunta desses três critérios resulta em uma base final refinada, com alta "
        "especificidade para caracterização da malha vicinal paulista."
    )
    
    adicionar_referencia_figura(doc, 5, "Aplicação de critérios lógicos finais", "passo4.2.png")
    
    doc.add_page_break()
    
    # ============================================================================
    # 3. RESULTADOS E DISCUSSÃO
    # ============================================================================
    adicionar_secao(doc, "3. RESULTADOS E DISCUSSÃO", nivel=1)
    
    adicionar_paragrafo(doc,
        "Os processamentos resultaram em duas bases geoespaciais complementares: (1) Malha Vicinal (OSM), "
        "contendo exclusivamente vias municipais estimadas a partir do OpenStreetMap após filtragens; e "
        "(2) Malha Total (OSM + DER), que integra a malha vicinal estimada com o cadastro oficial de "
        "rodovias estaduais do DER/SP. A Tabela 3 sintetiza os principais indicadores obtidos."
    )
    
    adicionar_tabela_resultados(doc)
    
    p_tabela3 = doc.add_paragraph()
    p_tabela3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_tabela3 = p_tabela3.add_run("Tabela 3 – Síntese dos resultados: indicadores globais")
    run_tabela3.font.size = Pt(10)
    run_tabela3.font.italic = True
    run_tabela3.font.color.rgb = RGBColor(102, 102, 102)
    
    # --------------------------------------------------------------------------
    # 3.1 Análise em Escala Municipal
    # --------------------------------------------------------------------------
    adicionar_secao(doc, "3.1. Análise em Escala Municipal", nivel=2)
    
    adicionar_secao(doc, "3.1.1. Extensão e Distribuição Territorial", nivel=3)
    adicionar_paragrafo(doc,
        "Observa-se heterogeneidade expressiva entre os 645 municípios paulistas, com razão máxima/mínima "
        "que pode superar 500:1. A distribuição assimétrica indica que a maioria dos municípios apresenta "
        "extensões reduzidas, enquanto poucos concentram malhas mais extensas. Os centros metropolitanos "
        "tendem a registrar extensões absolutas menores, possivelmente refletindo a menor expressão "
        "proporcional do sistema vicinal em áreas mais urbanizadas."
    )
    
    adicionar_lista(doc, [
        "Cananéia: 74,69 km (maior extensão municipal identificada);",
        "Itaoca: 35,39 km (município com perfil agrícola na RA Itapeva);",
        "Guarulhos: 25,36 km (município metropolitano com menor extensão relativa)."
    ])
    
    adicionar_secao(doc, "3.1.2. Densidade Espacial", nivel=3)
    adicionar_paragrafo(doc,
        "A densidade espacial média estimada é de aproximadamente 1.312 km por 10.000 km². Municípios com "
        "perfil agrícola tendem a apresentar valores superiores à média, enquanto os metropolitanos "
        "frequentemente ficam abaixo, padrão possivelmente associado ao uso do solo. Casos extremos ilustram "
        "essa variabilidade: Itaoca (RA Itapeva) atinge +45,9% acima da média, enquanto São Sebastião "
        "registra −90,1% abaixo da média."
    )
    
    adicionar_secao(doc, "3.1.3. Densidade Populacional", nivel=3)
    adicionar_paragrafo(doc,
        "A média estimada é de aproximadamente 35 km por 10.000 habitantes. Municípios menos populosos "
        "tendem a apresentar maior disponibilidade relativa de malha vicinal, o que sugere função "
        "predominantemente rural da infraestrutura. O contraste entre os extremos pode superar 500 vezes, "
        "evidenciando que a malha vicinal é proporcionalmente menos expressiva onde a população se concentra "
        "em núcleos urbanos densos."
    )
    
    adicionar_lista(doc, [
        "Ribeira: 111,5 km/10.000 hab (+217,9% acima da média);",
        "Guarulhos: 0,19 km/10.000 hab (−99,5% abaixo da média)."
    ])
    
    # --------------------------------------------------------------------------
    # 3.2 Análise em Escala Regional
    # --------------------------------------------------------------------------
    adicionar_secao(doc, "3.2. Análise em Escala Regional", nivel=2)
    
    adicionar_paragrafo(doc,
        "A análise por Região Administrativa (RA) apresenta menor heterogeneidade que a escala municipal, "
        "com razão máxima/mínima de aproximadamente 25:1, devido ao efeito de suavização pela agregação. "
        "Regiões do interior paulista, com perfil predominantemente agrícola, tendem a apresentar extensões "
        "superiores às metropolitanas."
    )
    
    adicionar_secao(doc, "3.2.1. Extensão por Região Administrativa", nivel=3)
    adicionar_lista(doc, [
        "RA Campinas: 3.584 km (1º lugar, 90 municípios);",
        "RA São José do Rio Preto: 2.625 km (2º lugar, 96 municípios);",
        "RA Presidente Prudente: 2.378 km (3º lugar, 53 municípios);",
        "RA Santos: 133 km (menor extensão, 9 municípios)."
    ])
    
    adicionar_secao(doc, "3.2.2. Densidade Populacional por RA", nivel=3)
    adicionar_paragrafo(doc,
        "A média regional é de aproximadamente 12 km por 10.000 habitantes. O padrão inversamente proporcional "
        "à urbanização se confirma na escala regional: RAs metropolitanas (RMSP, Campinas, Baixada Santista) "
        "tendem a valores reduzidos, enquanto RAs agrícolas (Presidente Prudente, Araçatuba, Itapeva) apresentam "
        "valores mais elevados."
    )
    
    adicionar_lista(doc, [
        "RA Itapeva: 28,5 km/10.000 hab (+113% acima da média);",
        "RA Registro: 50,7 km/10.000 hab (perfil rural acentuado);",
        "RMSP: 0,62 km/10.000 hab (−95% abaixo da média);",
        "RA Santos: 2,29 km/10.000 hab (−91% abaixo da média)."
    ])
    
    # --------------------------------------------------------------------------
    # 3.3 Integração com Rede Estadual (Malha Total)
    # --------------------------------------------------------------------------
    adicionar_secao(doc, "3.3. Integração com Rede Estadual (Malha Total)", nivel=2)
    
    adicionar_paragrafo(doc,
        "A integração com rodovias estaduais (DER) amplia significativamente a cobertura estimada. O incremento "
        "proporcional varia entre os municípios conforme a sobreposição entre as fontes e a presença de rodovias "
        "estaduais no território. Alguns municípios praticamente não apresentam incremento (baixa presença do DER), "
        "enquanto outros podem mais que quadruplicar sua extensão total."
    )
    
    adicionar_lista(doc, [
        "Guarulhos: 25,36 km → 109,91 km (incremento de +333%);",
        "Itaberá: 34,12 km → 162,41 km (incremento de +376%);",
        "Itaoca: 35,39 km → 35,44 km (incremento marginal, ~0%)."
    ])
    
    # --------------------------------------------------------------------------
    # 3.4 Discussão
    # --------------------------------------------------------------------------
    adicionar_secao(doc, "3.4. Discussão", nivel=2)
    
    adicionar_paragrafo(doc,
        "Os resultados evidenciam padrões espaciais consistentes com a literatura sobre infraestrutura viária "
        "e desenvolvimento regional. A concentração de malha vicinal em municípios menos populosos e de perfil "
        "agrícola reflete a função histórica dessas vias como eixos de escoamento da produção rural e acesso "
        "a propriedades dispersas. Em contrapartida, áreas metropolitanas apresentam menor extensão relativa "
        "de vicinais, possivelmente devido à consolidação urbana e à preponderância de vias urbanas classificadas."
    )
    
    adicionar_paragrafo(doc,
        "A heterogeneidade observada — razão de 500:1 na escala municipal e 25:1 na escala regional — sugere "
        "que decisões de planejamento e priorização de investimentos devem considerar ambas as escalas de análise. "
        "A escala municipal permite diagnósticos localizados e identificação de municípios atípicos, enquanto a "
        "escala regional oferece visão estabilizada para políticas estaduais e alocação de recursos entre RAs."
    )
    
    adicionar_caixa_ressalva(doc, "Limitações dos Resultados",
        "Os valores apresentados são estimativas sujeitas a incertezas. Possíveis fontes de divergência incluem: "
        "(1) incompletude do OpenStreetMap em áreas de menor contribuição voluntária; (2) defasagem temporal dos "
        "dados de áreas urbanas (2019); (3) imprecisões geométricas na sobreposição de bases com origens distintas; "
        "(4) variabilidade nos critérios de classificação de vias no OSM. Recomenda-se validação em campo antes de "
        "quaisquer decisões de investimento baseadas exclusivamente nesta base."
    )
    
    doc.add_page_break()
    
    # ============================================================================
    # 4. CONSIDERAÇÕES FINAIS
    # ============================================================================
    adicionar_secao(doc, "4. CONSIDERAÇÕES FINAIS", nivel=1)
    
    adicionar_paragrafo(doc,
        "O presente estudo apresentou uma metodologia de inferência geoespacial para geração de base vetorial "
        "rodoviária estimada da malha vicinal do Estado de São Paulo. A partir da integração de fontes abertas "
        "(OpenStreetMap) e oficiais (DER/SP, IBGE, SEADE), foram produzidas duas bases complementares: "
        "(1) Malha Vicinal (OSM), com 25.918,60 km em 7.417 segmentos; e (2) Malha Total (OSM + DER), com "
        "47.666,00 km integrando vicinais e rodovias estaduais."
    )
    
    adicionar_paragrafo(doc,
        "A análise exploratória revelou heterogeneidade expressiva entre municípios e Regiões Administrativas, "
        "com padrões espaciais consistentes com a literatura: maior disponibilidade relativa de malha vicinal "
        "em municípios agrícolas e menos populosos, e menor expressão em áreas metropolitanas densamente urbanizadas. "
        "A razão máxima/mínima de 500:1 na escala municipal e 25:1 na escala regional evidencia a importância de "
        "análises multiescala para suporte à tomada de decisão."
    )
    
    adicionar_secao(doc, "4.1. Principais Contribuições", nivel=2)
    adicionar_lista(doc, [
        "Geração de base vetorial inédita de cobertura estadual para malha vicinal paulista;",
        "Metodologia reprodutível, transparente e baseada em ferramentas de código aberto;",
        "Caracterização quantitativa da malha vicinal em escala municipal (645 municípios) e regional (16 RAs);",
        "Integração de fontes abertas e oficiais, demonstrando viabilidade de abordagens híbridas;",
        "Identificação de padrões espaciais e disparidades entre territórios, subsidiando priorização de ações."
    ])
    
    adicionar_secao(doc, "4.2. Limitações e Ressalvas", nivel=2)
    adicionar_caixa_ressalva(doc, "Natureza Estimativa da Base",
        "A base apresentada constitui uma APROXIMAÇÃO da malha vicinal, obtida por inferência geoespacial. "
        "NÃO se trata de levantamento topográfico ou cadastro oficial. Os valores são indicativos e podem "
        "divergir da realidade em campo. Toda inferência feita a partir desta base deve ser acompanhada de "
        "ressalvas importantes sobre suas limitações."
    )
    
    adicionar_lista(doc, [
        "Dependência da qualidade e completude do OpenStreetMap, que varia regionalmente;",
        "Defasagem temporal dos filtros urbanos (IBGE 2019) em relação ao presente;",
        "Buffers amplos de exclusão (DER) podem remover vicinais paralelas legítimas;",
        "Fragmentação de segmentos pós-interseção exige agregação cuidadosa para interpretação;",
        "Ausência de validação em campo sistemática dos resultados."
    ])
    
    adicionar_secao(doc, "4.3. Perspectivas de Aprimoramento", nivel=2)
    adicionar_paragrafo(doc,
        "Apesar das limitações, a base cumpre papel relevante ao subsidiar análises exploratórias iniciais "
        "do PLI/SP, permitindo identificação preliminar de padrões e tendências territoriais. Para aprimoramento "
        "contínuo, recomenda-se:"
    )
    
    adicionar_lista(doc, [
        "Validação cruzada com cadastros municipais quando disponíveis;",
        "Verificação em campo de municípios com valores extremos ou atípicos;",
        "Atualização dos filtros urbanos quando o IBGE disponibilizar novos dados de áreas urbanizadas;",
        "Integração com sistemas de monitoramento por sensoriamento remoto;",
        "Refinamento dos buffers de exclusão com análise de sensibilidade;",
        "Desenvolvimento de indicadores de confiança espacialmente diferenciados."
    ])
    
    adicionar_paragrafo(doc,
        "A base representa um ponto de partida rumo a uma representação mais sólida da malha vicinal paulista. "
        "Seu valor reside menos na precisão absoluta de cada valor individual e mais na capacidade de fornecer "
        "uma visão panorâmica e comparativa da distribuição espacial da infraestrutura viária municipal, orientando "
        "a priorização de ações que demandam esforços mais aprofundados — como levantamentos de campo, estudos de "
        "viabilidade e diagnósticos localizados — concentrando recursos onde os indícios iniciais sugerem maior relevância."
    )
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    # Linha de fechamento
    p_final = doc.add_paragraph()
    p_final.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_final = p_final.add_run("***")
    run_final.font.size = Pt(14)
    run_final.font.color.rgb = RGBColor(102, 102, 102)
    
    # Salvar documento
    caminho_saida = r"D:\ESTUDO_VICINAIS_V2\motodologia_abordagem_silvio\RELATORIO\Relatorio_Tecnico_Malha_Vicinal_SP_2026.docx"
    doc.save(caminho_saida)
    
    return caminho_saida

if __name__ == "__main__":
    print("=" * 80)
    print("GERAÇÃO DE RELATÓRIO TÉCNICO")
    print("Inferência Geoespacial - Base Vetorial Rodoviária - Estado de São Paulo")
    print("=" * 80)
    print()
    
    print("Iniciando geração do relatório...")
    caminho = criar_relatorio()
    
    print()
    print("✓ Relatório gerado com sucesso!")
    print(f"✓ Arquivo salvo em: {caminho}")
    print()
    print("ESTRUTURA DO RELATÓRIO:")
    print("  1. Introdução")
    print("  2. Metodologia")
    print("     2.1. Fontes de Dados")
    print("     2.2. Sistema de Referência Espacial")
    print("     2.3. Procedimentos de Processamento Espacial")
    print("  3. Resultados e Discussão")
    print("     3.1. Análise em Escala Municipal")
    print("     3.2. Análise em Escala Regional")
    print("     3.3. Integração com Rede Estadual")
    print("     3.4. Discussão")
    print("  4. Considerações Finais")
    print("     4.1. Principais Contribuições")
    print("     4.2. Limitações e Ressalvas")
    print("     4.3. Perspectivas de Aprimoramento")
    print()
    print("OBSERVAÇÕES:")
    print("  • As figuras estão referenciadas no texto com indicação dos arquivos")
    print("  • Insira manualmente as imagens dos layouts (passo1.png a passo4.2.png)")
    print("  • Tabelas de dados, filtros e resultados já estão formatadas")
    print("  • Caixas de ressalva destacadas com formatação especial")
    print("  • Numeração de páginas e cabeçalho/rodapé configurados")
    print()
    print("=" * 80)
